import os
import re
import io
import tempfile

import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import joblib
import shap
import numpy as np
import matplotlib.pyplot as plt
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches

from data_prep import choose_features

st.set_page_config(layout="wide", page_title="METABRIC Predictor (Classification + Regression)")

MODEL_DIR = "models"
CLASS_MODEL_FILE = "rf_pipeline.joblib"     # classification pipeline filename
REG_MODEL_FILE = "rf_reg_pipeline.joblib"   # regression pipeline filename

@st.cache_resource
def load_pipeline(fname):
    path = os.path.join(MODEL_DIR, fname)
    if not os.path.exists(path):
        return None
    return joblib.load(path)

st.title("🔬 METABRIC Predictor — Classification & Regression with SHAP")
st.markdown("Upload METABRIC clinical CSV or use the default file in `data/`")

uploaded = st.file_uploader("Upload CSV (optional)", type=["csv"])
if uploaded:
    df = pd.read_csv(uploaded)
else:
    default = "data/brca_metabric_clinical_data.csv"
    if os.path.exists(default):
        df = pd.read_csv(default)
    else:
        st.warning("No data uploaded and default data file not found.")
        st.stop()

st.write("Dataset preview")
st.dataframe(df.head())

task = st.sidebar.selectbox("Select task", ["Regression: predict survival months", "Classification: predict survival status"])

# Load pipeline according to task
if task.startswith("Regression"):
    pipe = load_pipeline(REG_MODEL_FILE)
    if pipe is None:
        st.error(f"Regression pipeline not found at {os.path.join(MODEL_DIR, REG_MODEL_FILE)}. Run train_model.py to create it.")
        st.stop()
else:
    pipe = load_pipeline(CLASS_MODEL_FILE)
    if pipe is None:
        st.error(f"Classification pipeline not found at {os.path.join(MODEL_DIR, CLASS_MODEL_FILE)}. Train classification pipeline or place it in models/.")
        st.stop()

# Prepare features
X_all = choose_features(df)
st.sidebar.header("Prediction options")
idx = st.sidebar.number_input("Patient row index", min_value=0, max_value=len(df)-1, value=0, step=1)
patient_row = df.iloc[[idx]]
patient_X = choose_features(patient_row)

# Get pipeline components
preprocessor = pipe.named_steps.get("preprocessor")
# find estimator step (non-preprocessor)
estimator = None
for name, step in pipe.named_steps.items():
    if name not in ("preprocessor", "selector", "scaler"):
        estimator = step

st.subheader("Prediction")
if task.startswith("Regression"):
    pred_months = float(pipe.predict(patient_X)[0])
    st.write(f"Estimated overall survival: **{pred_months:.1f} months**")
    target_value = pred_months
else:
    pred = pipe.predict(patient_X)[0]
    probs = pipe.predict_proba(patient_X)[0] if hasattr(pipe, "predict_proba") else None
    classes = getattr(estimator, "classes_", None)
    if classes is not None and set(classes) >= {0, 1}:
        label_map = {0: "Living", 1: "Deceased"}
        pred_label = label_map.get(pred, str(pred))
    else:
        pred_label = str(pred)
    st.write(f"Predicted label: **{pred_label}**")
    if probs is not None and classes is not None:
        try:
            idx_pos = list(classes).index(1)
            prob_pos = probs[idx_pos]
        except Exception:
            prob_pos = probs.max()
        st.write(f"Estimated probability (positive class): **{prob_pos:.3f}**")
    target_value = prob_pos if probs is not None else pred

st.write("Interpretation: model estimate only — not a clinical diagnosis.")

# SHAP explanation
st.subheader("Local SHAP explanation")
if preprocessor is None or estimator is None:
    st.warning("Pipeline missing expected components. SHAP cannot be computed.")
else:
    # Transform input
    X_trans = preprocessor.transform(patient_X)

    # Try TreeExplainer first, fallback to generic Explainer
    try:
        explainer = shap.TreeExplainer(estimator)
    except Exception:
        try:
            # try using background from train set (use subset of X_all)
            bg = preprocessor.transform(choose_features(df).sample(min(100, len(df)), random_state=0))
            explainer = shap.Explainer(estimator, bg)
        except Exception:
            explainer = shap.Explainer(estimator)

    shap_vals = explainer.shap_values(X_trans)

    # feature names
    try:
        feature_names = preprocessor.get_feature_names_out()
    except Exception:
        feature_names = []
        for name_tr, trans, cols in preprocessor.transformers_:
            if hasattr(trans, "get_feature_names_out"):
                try:
                    feats = trans.get_feature_names_out(cols)
                except Exception:
                    feats = trans.get_feature_names_out()
            elif hasattr(trans, "get_feature_names"):
                try:
                    feats = trans.get_feature_names(cols)
                except Exception:
                    feats = trans.get_feature_names()
            else:
                feats = cols
            feature_names.extend(list(feats))

    # Robust extractor to produce 1-D local shap vector and base value
    def _extract_local_shap(shap_vals, feature_names, task_mode):
        n_feat = len(feature_names)
        # handle list (common for multiclass)
        if isinstance(shap_vals, list):
            # classification: prefer positive class (index 1) when available
            if task_mode.startswith("Classification") and len(shap_vals) > 1:
                arr = np.asarray(shap_vals[1])
            else:
                arr = np.asarray(shap_vals[0])
        else:
            arr = np.asarray(shap_vals)

        # If arr already 1-D and matches features
        if arr.ndim == 1 and arr.size == n_feat:
            return arr, getattr(explainer, "expected_value", None)

        # Find axis that equals n_feat
        axes = [i for i, s in enumerate(arr.shape) if s == n_feat]
        if axes:
            feat_axis = axes[-1]
            arrm = np.moveaxis(arr, feat_axis, -1)  # features on last axis
            sel = arrm
            while sel.ndim > 1:
                sel = sel[0]
            return np.asarray(sel), getattr(explainer, "expected_value", None)

        # If arr has shape (n_samples, n_feat) or (n_samples, n_classes, n_feat)
        if arr.ndim >= 2:
            # try common case: first axis is samples
            if arr.shape[0] >= 1 and arr.shape[-1] == n_feat:
                sel = arr[0]
                if sel.ndim == 2 and sel.shape[-1] == n_feat:
                    if task_mode.startswith("Classification") and sel.shape[0] > 1:
                        return np.asarray(sel[1]), getattr(explainer, "expected_value", None)
                    else:
                        return np.asarray(sel[0]), getattr(explainer, "expected_value", None)
                elif sel.ndim == 1 and sel.size == n_feat:
                    return np.asarray(sel), getattr(explainer, "expected_value", None)

        # fallback: flatten and take first n_feat entries if possible
        flat = arr.ravel()
        if flat.size >= n_feat:
            return flat[:n_feat], getattr(explainer, "expected_value", None)

        raise ValueError(f"Cannot convert shap array with shape {arr.shape} to a 1-D vector of length {n_feat}")

    try:
        sv_local, base_value = _extract_local_shap(shap_vals, feature_names, task)
    except Exception as e:
        st.error(f"Failed to parse SHAP values: {e}")
        raise

    # If expected_value is array, choose index matching class if classification
    if isinstance(base_value, (list, tuple, np.ndarray)):
        if task.startswith("Classification") and len(base_value) > 1:
            base_value = base_value[1]
        else:
            base_value = base_value[0]

    # verify dimensions
    if sv_local.ndim != 1 or sv_local.shape[0] != len(feature_names):
        st.error(f"Extracted SHAP vector shape {sv_local.shape} does not match feature names length {len(feature_names)}")
        raise ValueError("SHAP extraction mismatch")

    contribs = pd.Series(sv_local, index=feature_names)
    st.write("Top positive (increase) and negative (decrease) contributors:")
    st.write(pd.concat([contribs.sort_values(ascending=False).head(10), contribs.sort_values().head(10)]))

    # Interactive force plot (HTML)
    try:
        force_html = shap.plots.force(base_value, sv_local, feature_names=feature_names, matplotlib=False)
        shap.save_html("force_local.html", force_html)
        with open("force_local.html", "r", encoding="utf-8") as f:
            components.html(f.read(), height=150, scrolling=True)
    except Exception:
        st.info("Interactive SHAP force plot not available in this environment.")

    # Static PNG for report
    png_path = None
    try:
        shap.force_plot(base_value, sv_local, feature_names=feature_names, matplotlib=True, show=False)
        png_path = os.path.join(tempfile.gettempdir(), "force_local.png")
        plt.savefig(png_path, bbox_inches="tight", dpi=150)
        plt.close()
    except Exception:
        png_path = None

    # Build PDF report
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("METABRIC Prediction Report", styles["Title"]))
    story.append(Spacer(1, 12))
    if task.startswith("Regression"):
        story.append(Paragraph(f"Regression prediction (overall survival months): <b>{target_value:.1f}</b>", styles["Normal"]))
    else:
        story.append(Paragraph(f"Classification prediction: <b>{pred_label}</b>", styles["Normal"]))
        if probs is not None:
            story.append(Paragraph(f"Predicted probability (positive class): <b>{prob_pos:.3f}</b>", styles["Normal"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Top feature contributors (local):", styles["Heading2"]))
    top_pos = contribs.sort_values(ascending=False).head(10)
    for feat, val in top_pos.items():
        story.append(Paragraph(f"{feat}: {val:.4f}", styles["Normal"]))
    story.append(Spacer(1, 12))
    if png_path and os.path.exists(png_path):
        story.append(Image(png_path, width=500, height=150))
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        doc = SimpleDocTemplate(tmp_pdf.name)
        doc.build(story)
        tmp_pdf.seek(0)
        pdf_bytes = tmp_pdf.read()
    st.download_button("Download PDF report", data=pdf_bytes, file_name="metabric_prediction_report.pdf", mime="application/pdf")

    # Prepare DOCX plain-language report
    FEATURE_NAME_MAP = {
        "num__TumorSize": "Tumor size",
        "num__TumorStage": "Tumor stage",
        "Nottingham prognostic index": "Nottingham prognostic index",
        "Integrative Cluster": "Integrative cluster",
        "num__Relapse Free Status (Months)": "Relapse-free time (months)",
        "cat__Relapse Free Status_0:Not_Recurred": "No recurrence recorded",
        "cat__Relapse Free Status_1:Recurred": "Had recurrence",
    }

    def pretty_name(feat: str) -> str:
        if feat in FEATURE_NAME_MAP:
            return FEATURE_NAME_MAP[feat]
        feat_clean = re.sub(r'^(num__|cat__|onehot__|x0__)', '', feat)
        feat_clean = feat_clean.replace(":", " = ")
        feat_clean = re.sub(r'[_]+', ' ', feat_clean)
        feat_clean = re.sub(r'\s+', ' ', feat_clean).strip()
        if len(feat_clean.split()) <= 4:
            return feat_clean.title()
        return feat_clean

    def sentence_for(feat: str, val: float) -> str:
        pretty = pretty_name(feat)
        amt = round(abs(float(val)), 1)
        if val > 0:
            return f"- {pretty}: increased the prediction by about {amt} (model units)."
        else:
            return f"- {pretty}: decreased the prediction by about {amt} (model units)."

    def create_word_report(pred_val, top_contribs, png):
        doc = Document()
        doc.add_heading("METABRIC Prediction Report", level=1)
        if task.startswith("Regression"):
            doc.add_paragraph(f"Estimated overall survival: {pred_val:.1f} months.")
        else:
            doc.add_paragraph(f"Predicted label: {pred_label}")
            if probs is not None:
                doc.add_paragraph(f"Predicted probability (positive class): {prob_pos:.3f}")
        doc.add_paragraph("Important: this is a model estimate, not a clinical diagnosis.")
        doc.add_heading("Top factors and plain-language explanation", level=2)
        shown = top_contribs.copy()
        if shown.shape[0] > 10:
            shown = shown.head(10)
        for feat, val in shown.items():
            doc.add_paragraph(sentence_for(feat, val))
        doc.add_heading("Feature mapping (technical → friendly)", level=2)
        for feat in shown.index:
            doc.add_paragraph(f"- {feat} → {pretty_name(feat)}")
        if png and os.path.exists(png):
            try:
                doc.add_page_break()
                doc.add_heading("SHAP visual explanation", level=2)
                doc.add_picture(png, width=Inches(6))
            except Exception:
                pass
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.read()

    word_bytes = create_word_report(target_value, contribs.sort_values(ascending=False), png_path)
    st.download_button("Download Word explanation (DOCX)", data=word_bytes, file_name="metabric_prediction_explanation.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
