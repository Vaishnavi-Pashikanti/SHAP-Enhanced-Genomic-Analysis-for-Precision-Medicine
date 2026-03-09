import streamlit as st
import pandas as pd
import joblib
import shap
import numpy as np
import os
import re
import streamlit.components.v1 as components
import matplotlib.pyplot as plt
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
import tempfile
import io
from docx import Document
from docx.shared import Inches

MODEL_PATH = os.path.join("models", "rf_reg_pipeline.joblib")

@st.cache_resource
def load_pipeline(path=MODEL_PATH):
    if not os.path.exists(path):
        st.error(f"Regression pipeline not found at {path}. Run train_model.py first.")
        st.stop()
    return joblib.load(path)

st.title("🔬 METABRIC Survival Months Predictor (Regression)")

uploaded = st.file_uploader("Upload METABRIC clinical CSV (optional)", type=["csv"])
if uploaded:
    df = pd.read_csv(uploaded)
else:
    default = "data/brca_metabric_clinical_data.csv"
    if os.path.exists(default):
        df = pd.read_csv(default)
    else:
        st.warning("No data uploaded and default data file not found.")
        st.stop()

st.write("Dataset preview")
st.dataframe(df.head())

from data_prep import choose_features

pipe = load_pipeline()
X = choose_features(df)

st.sidebar.header("Prediction options")
idx = st.sidebar.number_input("Patient row index", min_value=0, max_value=len(df)-1, value=0, step=1)
patient_row = df.iloc[[idx]]
patient_X = choose_features(patient_row)

# Regression prediction (months)
pred_months = float(pipe.predict(patient_X)[0])
st.subheader("Predicted overall survival (months)")
st.write(f"Estimated survival: **{pred_months:.1f} months**")
st.write("Note: model estimate only — not a clinical diagnosis.")

# SHAP explanation for regressor
st.subheader("Local SHAP explanation")
# Use regressor object inside pipeline
reg = pipe.named_steps.get("reg", pipe.named_steps.get("clf", None))
preprocessor = pipe.named_steps.get("preprocessor", None)

if reg is None or preprocessor is None:
    st.warning("Pipeline missing expected steps ('preprocessor' or 'reg'). Cannot compute SHAP.")
else:
    # transform features for model
    X_trans = preprocessor.transform(patient_X)
    # TreeExplainer for tree-based regressors
    explainer = shap.TreeExplainer(reg)
    shap_vals = explainer.shap_values(X_trans)
    # get feature names after preprocessing
    try:
        feature_names = preprocessor.get_feature_names_out()
    except Exception:
        feature_names = []
        for name, trans, cols in preprocessor.transformers_:
            if hasattr(trans, "get_feature_names_out"):
                feats = trans.get_feature_names_out(cols)
            elif hasattr(trans, "get_feature_names"):
                feats = trans.get_feature_names(cols)
            else:
                feats = cols
            feature_names.extend(list(feats))

    # shap_vals may be array
    sv = shap_vals if not isinstance(shap_vals, list) else shap_vals[0]
    contribs = pd.Series(sv[0], index=feature_names)
    st.write(contribs.sort_values(ascending=False).head(10))

    # Force plot (HTML) and PNG for PDF
    expected_value = explainer.expected_value
    # shap.force_plot returns matplotlib/HTML depending on args
    force_html = shap.plots.force(expected_value, sv[0], feature_names=feature_names, matplotlib=False)
    shap.save_html("force_reg.html", force_html)
    with open("force_reg.html", "r", encoding="utf-8") as f:
        components.html(f.read(), height=200, scrolling=True)

    # Static PNG for report
    shap.force_plot(expected_value, sv[0], feature_names=feature_names, matplotlib=True, show=False)
    png_path = os.path.join(tempfile.gettempdir(), "force_reg.png")
    plt.savefig(png_path, bbox_inches="tight", dpi=150)
    plt.close()

    # PDF report
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("METABRIC Regression Prediction Report", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Predicted overall survival (months): <b>{pred_months:.1f}</b>", styles["Normal"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Top feature contributions:", styles["Heading2"]))
    top = contribs.sort_values(ascending=False).head(10)
    for feat, val in top.items():
        story.append(Paragraph(f"{feat}: {val:.4f}", styles["Normal"]))
    story.append(Spacer(1, 12))
    story.append(Image(png_path, width=500, height=200))

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        doc = SimpleDocTemplate(tmp_pdf.name)
        doc.build(story)
        tmp_pdf.seek(0)
        pdf_bytes = tmp_pdf.read()

    st.download_button("Download PDF report", data=pdf_bytes, file_name="metabric_regression_report.pdf", mime="application/pdf")

# ------------------
    # mapping for known technical features -> friendly labels
    FEATURE_NAME_MAP = {
        "num__TumorSize": "Tumor size",
        "num__TumorStage": "Tumor stage",
        "Nottingham prognostic index": "Nottingham prognostic index",
        "Integrative Cluster": "Integrative cluster",
        "num__Relapse Free Status (Months)": "Relapse-free time (months)",
        "cat__Relapse Free Status_0:Not_Recurred": "No recurrence recorded",
        "cat__Relapse Free Status_1:Recurred": "Had recurrence",
        # add more known mappings here as needed
    }

    def pretty_name(feat: str) -> str:
        # exact map first
        if feat in FEATURE_NAME_MAP:
            return FEATURE_NAME_MAP[feat]
        # handle patterns like cat__Feature_1:Value or cat__Feature:Value
        feat_clean = feat
        # remove common prefixes
        feat_clean = re.sub(r'^(num__|cat__|onehot__|x0__)', '', feat_clean)
        # replace colon separators used by some one-hot encoders
        feat_clean = feat_clean.replace(":", " = ")
        # replace underscores with spaces and collapse multiple spaces
        feat_clean = re.sub(r'[_]+', ' ', feat_clean)
        feat_clean = re.sub(r'\s+', ' ', feat_clean).strip()
        # title-case short phrases, keep acronyms upper-case
        if len(feat_clean.split()) <= 4:
            return feat_clean.title()
        return feat_clean

    def sentence_for(feat: str, val: float) -> str:
        pretty = pretty_name(feat)
        amt = round(abs(float(val)), 1)
        if val > 0:
            # feature pushed prediction up
            if "no recurrence" in pretty.lower() or "not recurred" in pretty.lower():
                # special phrasing for "no recurrence"
                return f"- {pretty}: associated with a higher estimate, increasing predicted survival by about {amt} months."
            if "=" in pretty:  # categorical like "Relapse Free Status = Recurred"
                return f"- Having {pretty.split('=')[-1].strip()} {pretty.split('=')[0].strip().lower()} was associated with an increase of about {amt} months."
            return f"- Higher {pretty} increased the estimated survival by about {amt} months."
        else:
            # feature pushed prediction down
            if "recurr" in pretty.lower():
                return f"- {pretty}: associated with a lower estimate, decreasing predicted survival by about {amt} months."
            if "=" in pretty:
                return f"- Having {pretty.split('=')[-1].strip()} {pretty.split('=')[0].strip().lower()} was associated with a decrease of about {amt} months."
            return f"- Higher {pretty} decreased the estimated survival by about {amt} months."

    def create_word_report(pred_months: float, top_contribs: pd.Series, png_path: str) -> bytes:
        """
        Build a Word document with plain-language sentences derived from SHAP contributions.
        top_contribs: pd.Series indexed by technical feature name, values are SHAP contributions (months).
        """
        doc = Document()
        doc.add_heading("METABRIC Regression Prediction Report", level=1)
        doc.add_paragraph(f"Estimated overall survival: {pred_months:.1f} months.")
        doc.add_paragraph(
            "Important: this is a statistical model estimate, not a medical diagnosis. "
            "The list below explains which patient factors moved the estimate up or down and by approximately how many months."
        )
        doc.add_heading("Top factors and plain‑language explanation", level=2)

        # use top_contribs (already sorted descending) but show absolute-important ordering
        # show up to 10 items
        shown = top_contribs.copy()
        if shown.shape[0] > 10:
            shown = shown.head(10)

        # produce exact sentences
        for feat, val in shown.items():
            sent = sentence_for(feat, val)
            doc.add_paragraph(sent)

        # Add a short glossary mapping technical names to friendly labels
        doc.add_heading("Feature name mapping (technical → friendly)", level=2)
        for feat in shown.index:
            doc.add_paragraph(f"- {feat} → {pretty_name(feat)}")

        # Add SHAP figure if available
        if os.path.exists(png_path):
            try:
                doc.add_page_break()
                doc.add_heading("Visual explanation (SHAP)", level=2)
                doc.add_picture(png_path, width=Inches(6))
            except Exception:
                pass

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.read()

    # Prepare Word bytes and provide download button
    # use 'top' Series computed earlier (contribs.sort_values(...).head(10))
    word_bytes = create_word_report(pred_months, top, png_path)
    st.download_button(
        "Download Word explanation (DOCX)",
        data=word_bytes,
        file_name="metabric_regression_explanation.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )